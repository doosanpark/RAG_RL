"""report_solution_a.md -> .docx 변환 (헤더/표/리스트/볼드/코드/이미지).

실행:
    .\.venv\Scripts\python.exe -u -X utf8 -m src.sol_a.md_to_docx
"""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor, Inches

ROOT = Path(__file__).resolve().parent.parent.parent
SRC = ROOT / "report_solution_a.md"
OUT = ROOT / "report_solution_a.docx"
FIG = ROOT / "results/sol_a_learning_curves.png"
KFONT = "맑은 고딕"  # Malgun Gothic


def set_korean_font(doc: Document) -> None:
    style = doc.styles["Normal"]
    style.font.name = KFONT
    style.font.size = Pt(10.5)
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    for attr in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
        rfonts.set(qn(attr), KFONT)


def add_runs(paragraph, text: str) -> None:
    """**bold**, `code` 인라인 파싱."""
    # 토큰 분리: **...** 또는 `...`
    parts = re.split(r"(\*\*.+?\*\*|`.+?`)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            r = paragraph.add_run(part[2:-2]); r.bold = True
        elif part.startswith("`") and part.endswith("`"):
            r = paragraph.add_run(part[1:-1]); r.font.name = "Consolas"
            r.font.color.rgb = RGBColor(0xB0, 0x30, 0x30)
        else:
            paragraph.add_run(part)


def add_table(doc: Document, rows: list[list[str]]) -> None:
    header, body = rows[0], rows[1:]
    t = doc.add_table(rows=1, cols=len(header))
    t.style = "Light Grid Accent 1"
    for j, cell in enumerate(header):
        p = t.rows[0].cells[j].paragraphs[0]
        add_runs(p, cell.strip())
        for run in p.runs:
            run.bold = True
    for row in body:
        cells = t.add_row().cells
        for j, cell in enumerate(row[: len(header)]):
            add_runs(cells[j].paragraphs[0], cell.strip())


def main() -> None:
    lines = SRC.read_text(encoding="utf-8").splitlines()
    doc = Document()
    set_korean_font(doc)

    i = 0
    in_code = False
    code_buf: list[str] = []
    while i < len(lines):
        line = lines[i]

        # 코드펜스
        if line.strip().startswith("```"):
            if in_code:
                p = doc.add_paragraph()
                r = p.add_run("\n".join(code_buf))
                r.font.name = "Consolas"; r.font.size = Pt(9)
                code_buf = []
                in_code = False
            else:
                in_code = True
            i += 1
            continue
        if in_code:
            code_buf.append(line); i += 1; continue

        # 그림 삽입 지점: "## 4. 분석" 직전
        if line.startswith("## 4.") and FIG.exists():
            cap = doc.add_paragraph()
            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            doc.add_picture(str(FIG), width=Inches(5.5))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
            fc = doc.add_paragraph(); fc.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = fc.add_run("그림 1. 3-seed GRPO dev F1 학습곡선 (improve→peak→drift, dev-best로 peak 보존)")
            r.italic = True; r.font.size = Pt(9)

        # 표 (| 로 시작하고 다음 줄이 |---| )
        if line.lstrip().startswith("|") and i + 1 < len(lines) and re.match(r"^\s*\|[\s:|-]+\|\s*$", lines[i + 1]):
            tbl_lines = [line]
            j = i + 2
            while j < len(lines) and lines[j].lstrip().startswith("|"):
                tbl_lines.append(lines[j]); j += 1
            rows = []
            for k, tl in enumerate(tbl_lines):
                if k == 1:
                    continue  # separator
                cells = [c for c in tl.strip().strip("|").split("|")]
                rows.append(cells)
            add_table(doc, rows)
            i = j
            continue

        # 헤더
        if line.startswith("#"):
            m = re.match(r"^(#+)\s*(.*)$", line)
            level = min(len(m.group(1)), 4)
            doc.add_heading(m.group(2).strip(), level=level - 0 if level > 1 else 0)
            i += 1
            continue

        # 수평선
        if line.strip() in ("---", "***", "___"):
            i += 1
            continue

        # 불릿 (- 또는 숫자.)
        mbul = re.match(r"^(\s*)[-*]\s+(.*)$", line)
        mnum = re.match(r"^(\s*)(\d+)\.\s+(.*)$", line)
        if mbul:
            indent = len(mbul.group(1)) // 2
            p = doc.add_paragraph(style="List Bullet")
            if indent:
                p.paragraph_format.left_indent = Inches(0.25 * (indent + 1))
            add_runs(p, mbul.group(2))
            i += 1
            continue
        if mnum:
            p = doc.add_paragraph(style="List Number")
            add_runs(p, mnum.group(3))
            i += 1
            continue

        # 빈 줄
        if not line.strip():
            i += 1
            continue

        # 일반 문단 (인용 > 포함)
        text = line[2:] if line.startswith("> ") else line
        p = doc.add_paragraph()
        add_runs(p, text)
        i += 1

    doc.save(str(OUT))
    print(f"[saved] {OUT}")


if __name__ == "__main__":
    main()
